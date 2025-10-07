from ..services import services
from ... import utils as ut
import os
import subprocess
import uuid
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import logging
import random

log = logging.getLogger(__name__)

SIMULATE_MANUAL_REVIEW = os.environ.get("SIMULATE_MANUAL_REVIEW", False)


def create_pd_pdf(data: dict):
    """
    Creates a Position Description PDF, robustly handling various data types.
    It processes strings, lists, and dictionaries. If a value expected to be a
    dictionary (like 'qualifications') is a string, it uses the string directly.
    """
    # --- Filename and Document Setup ---
    job_title = data.get("job_title")
    position_number = data.get("position_number")

    identifier = None

    if isinstance(position_number, str) and position_number.strip():
        identifier = position_number.strip().replace(" ", "_")
    else:
        identifier = str(uuid.uuid4())

    output_dir = "data/source/pds"
    os.makedirs(output_dir, exist_ok=True)
    docx_filename = os.path.join(output_dir, f"{identifier}.docx")
    pdf_filename = os.path.join(output_dir, f"{identifier}.pdf")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # --- Populate Document ---

    # Header
    company = data.get("company")
    if isinstance(company, str) and company.strip():
        doc.add_heading(company.strip(), level=0)
    if isinstance(job_title, str) and job_title.strip():
        doc.add_heading(job_title.strip(), level=1)
    doc.add_paragraph()

    # Key Details Section
    details = {
        "Department": data.get("department"),
        "Reports To": data.get("reports_to"),
    }
    for label, value in details.items():
        if isinstance(value, str) and value.strip():
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value.strip())

    # Position Summary
    summary = data.get("summary")
    if isinstance(summary, str) and summary.strip():
        doc.add_heading("Position Summary", level=2)
        doc.add_paragraph(summary.strip())

    # Key Responsibilities
    responsibilities = data.get("responsibilities")
    if isinstance(responsibilities, list) and responsibilities:
        # Filter for non-empty strings before adding the heading
        valid_items = [
            str(item).strip()
            for item in responsibilities
            if isinstance(item, str) and str(item).strip()
        ]
        if valid_items:
            doc.add_heading("Key Responsibilities", level=2)
            for item in valid_items:
                doc.add_paragraph(item, style="List Bullet")

    # Qualifications (Handles both dict and str)
    qual_data = data.get("qualifications")
    if isinstance(qual_data, dict):
        qual_content = []
        qual_fields = ["Education", "Experience", "Certifications"]
        for field in qual_fields:
            value = qual_data.get(field.lower())
            if isinstance(value, str) and value.strip():
                qual_content.append((field, value.strip()))

        if qual_content:
            doc.add_heading("Qualifications & Experience", level=2)
            for label, text in qual_content:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(text)
    elif isinstance(qual_data, str):
        doc.add_heading("Qualifications & Experience", level=2)
        doc.add_paragraph(qual_data.strip())

    # Skills (Handles both dict and str)
    skills_data = data.get("skills")
    if isinstance(skills_data, dict):
        hard_skills = skills_data.get("hard_skills")
        soft_skills = skills_data.get("soft_skills")

        valid_hard = isinstance(hard_skills, list) and any(
            isinstance(s, str) and s.strip() for s in hard_skills
        )
        valid_soft = isinstance(soft_skills, list) and any(
            isinstance(s, str) and s.strip() for s in soft_skills
        )

        if valid_hard or valid_soft:
            doc.add_heading("Skills", level=2)
            if valid_hard:
                doc.add_paragraph().add_run("Hard Skills").bold = True
                for skill in hard_skills:
                    if isinstance(skill, str) and skill.strip():
                        doc.add_paragraph(skill.strip(), style="List Bullet")
            if valid_soft:
                doc.add_paragraph().add_run("Soft Skills").bold = True
                for skill in soft_skills:
                    if isinstance(skill, str) and skill.strip():
                        doc.add_paragraph(skill.strip(), style="List Bullet")
    elif isinstance(skills_data, str):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(skills_data.strip())

    # --- Save Word Document and Convert to PDF ---
    doc.save(docx_filename)
    log.info(f"✅ Created Word document: {docx_filename}")

    log.info("⏳ Starting PDF conversion with LibreOffice...")
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_filename,
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    log.info(f"✅ Converted to PDF: {pdf_filename}")
    if os.path.exists(docx_filename):
        os.remove(docx_filename)
        log.info(f"🗑️ Removed temporary file: {docx_filename}")


def create_cv_pdf(data: dict, folder: str, application_id: str):
    """
    Creates a CV PDF from a dictionary, robustly handling missing data.

    This function generates a .docx file using the provided dictionary, formats it
    into a clean CV layout, and then converts it to a PDF using LibreOffice.
    The temporary .docx file is removed after the conversion.

    Args:
        data (dict): A dictionary containing all the CV information.
        folder (str): Folder to save pdf in.
        application_id (str): Unique identifier for the application.

    Returns:
        str: The file path to the generated PDF, or None if creation failed.
    """
    # --- Filename and Document Setup ---
    personal_info = data.get("personal_info", {})

    output_dir = f"data/source/cvs/{folder}"
    os.makedirs(output_dir, exist_ok=True)
    docx_filename = os.path.join(output_dir, f"{application_id}.docx")
    pdf_filename = os.path.join(output_dir, f"{application_id}.pdf")

    doc = Document()
    # Set document margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # --- 1. Header (Personal Info) ---
    if personal_info:
        # Full Name
        if personal_info.get("full_name"):
            heading = doc.add_heading(personal_info["full_name"], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Job Title
        if personal_info.get("job_title"):
            p = doc.add_paragraph(personal_info["job_title"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact Details
        contact_items = [
            personal_info.get("email"),
            personal_info.get("phone"),
            personal_info.get("linkedin"),
            personal_info.get("github"),
            personal_info.get("website"),
            personal_info.get("address"),
        ]
        # Filter out None or empty strings
        contact_line = " | ".join(
            item for item in contact_items if isinstance(item, str) and item.strip()
        )
        if contact_line:
            p = doc.add_paragraph(contact_line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 2. Summary ---
    summary = personal_info.get("summary")
    if isinstance(summary, str) and summary.strip():
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary)

    # --- 3. Work Experience ---
    experience = data.get("work_experience")
    if isinstance(experience, list) and experience:
        doc.add_heading("Work Experience", level=2)
        for job in experience:
            if isinstance(job, dict):
                p = doc.add_paragraph()
                p.add_run(job.get("job_title", "Job Title")).bold = True
                p.add_run(f" at {job.get('company', 'Company')}")

                date_line = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
                location = job.get("location", "")
                # Create a right-aligned tab stop for the date
                p = doc.add_paragraph()
                p.add_run(location)
                p.add_run("\t")
                p.add_run(date_line)
                p.paragraph_format.tab_stops.add_tab_stop(
                    Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT
                )

                responsibilities = job.get("responsibilities")
                if isinstance(responsibilities, list):
                    for item in responsibilities:
                        if isinstance(item, str) and item.strip():
                            doc.add_paragraph(item, style="List Bullet")
                doc.add_paragraph()  # Add a space after the entry

    # --- 4. Education ---
    education = data.get("education")
    if isinstance(education, list) and education:
        doc.add_heading("Education", level=2)
        for edu in education:
            if isinstance(edu, dict):
                p = doc.add_paragraph()
                p.add_run(edu.get("degree", "Degree")).bold = True
                p.add_run(f", {edu.get('institution', 'Institution')}")

                grad_line = f"Graduated: {edu.get('graduation_year', '')}"
                location = edu.get("location", "")
                p = doc.add_paragraph()
                p.add_run(location)
                p.add_run("\t")
                p.add_run(grad_line)
                p.paragraph_format.tab_stops.add_tab_stop(
                    Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT
                )

                if edu.get("details"):
                    doc.add_paragraph(edu.get("details"), style="List Bullet")

    # --- 5. Skills ---
    skills = data.get("skills")
    if isinstance(skills, dict) and skills:
        doc.add_heading("Skills", level=2)
        for category, skill_list in skills.items():
            if isinstance(skill_list, list) and skill_list:
                # Format category title (e.g., "programming_languages" -> "Programming Languages")
                category_title = category.replace("_", " ").title()
                p = doc.add_paragraph()
                p.add_run(category_title).bold = True

                # Special handling for list of dictionaries like 'languages'
                if all(isinstance(s, dict) for s in skill_list):
                    items = [
                        f"{s.get('language', '')} ({s.get('proficiency', '')})"
                        for s in skill_list
                    ]
                else:  # Assumes list of strings
                    items = [str(s) for s in skill_list]

                doc.add_paragraph(", ".join(items))

    # --- 6. Projects ---
    projects = data.get("projects")
    if isinstance(projects, list) and projects:
        doc.add_heading("Projects", level=2)
        for project in projects:
            if isinstance(project, dict):
                p = doc.add_paragraph()
                p.add_run(project.get("name", "Project Name")).bold = True
                if project.get("link"):
                    p.add_run(f" - {project.get('link')}")
                if project.get("description"):
                    doc.add_paragraph(project.get("description"))
                if project.get("technologies"):
                    tech_list = project.get("technologies", [])
                    tech_str = ", ".join(
                        tech for tech in tech_list if isinstance(tech, str)
                    )
                    doc.add_paragraph(f"Technologies: {tech_str}")

    # --- Save Word Document and Convert to PDF ---
    doc.save(docx_filename)
    log.info(f"✅ Created Word document: {docx_filename}")

    log.info("⏳ Starting PDF conversion with LibreOffice...")
    # Note: On macOS, you might need the full path to the soffice binary
    # e.g., /Applications/LibreOffice.app/Contents/MacOS/soffice
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_filename,
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    if os.path.exists(docx_filename):
        os.remove(docx_filename)
        log.info(f"🗑️  Removed temporary file: {docx_filename}")
    log.info(f"✅ Converted to PDF: {pdf_filename}")

    return pdf_filename


def generate_demographic_profiles(position_number: str, application_id: str):
    """
    Generates a random demographic profile.

    Args:
        position_number (str): Identifier for position.
        application_id (str): Identifier for application.

    Returns:
        A dictionary representing a demographic profile.
    """
    # Define the possible options for each demographic characteristic
    atsi_options = [
        "Yes, Aboriginal",
        "Yes, Torres Strait Islander",
        "Yes, both Aboriginal and Torres Strait Islander",
        "No",
        "Prefer not to say",
    ]
    disability_options = ["Yes", "No", "Prefer not to say"]
    age_options = [
        "Under 18",
        "18-24",
        "25-34",
        "35-44",
        "45-54",
        "55-64",
        "65+",
        "Prefer not to say",
    ]
    gender_options = [
        "Woman",
        "Man",
        "Non-binary",
        "I use a different term",
        "Prefer not to say",
    ]
    lgbtqia_options = ["Yes", "No", "Prefer not to say"]
    carer_options = ["Yes", "No", "Prefer not to say"]
    language_options = ["Yes", "No", "Prefer not to say"]

    profile = {
        "position_number": position_number,
        "application_id": application_id,
        "atsi_identity": random.choice(atsi_options),
        "disability_status": random.choice(disability_options),
        "age_range": random.choice(age_options),
        "gender_identity": random.choice(gender_options),
        "lgbtqia_community": random.choice(lgbtqia_options),
        "carer_responsibilities": random.choice(carer_options),
        "speaks_language_other_than_english": random.choice(language_options),
    }

    return profile


if __name__ == "__main__":

    response_pd = services.llm_hot.invoke(
        """
        ### BACKGROUND ###
        You are an expert HR manager. Your primary function is to create job position descriptions.

        ### TASK ###
        Create the contents of a jsonl file.  
        Each line corresponds to the information to be contained in a position description document.
        The entries in each line are
            - position_number: a unique identifier for the position
            - job_title: the title of the position
            - level: a number 1,2,3,4,5, or 6 which indicates the seniority of the position. Must be included.
            - department: which part of the business the position works in (e.g. "Claims Management")
            - reports_to: the position this position reports to (e.g. "Claims Manager")
            - summary: a brief description of the position (e.g. "To efficiently and empathetically manage a
            portfolio of insurance claims from notification to settlement, ensuring a high level of customer
            satisfaction and adherence to company policies and regulatory requirements.")
            - responsibilities: a list of things the position will need to do (e.g. contains things like "Assess
            and process new insurance claims in a timely and accurate manner.")
            - qualifications: Information on education, experience, certifications.
            - skills: practical skills relevant to the position.

        YOU MUST PROVIDE THE `qualifications` AND `skills` ENTRIES AS STRINGS.
        
        The positions all correspond to jobs at Van Gogh Museum. 

        Create information for five distinct positions, each with it's own line in the jsonl.

        Two of the positions should be very similar to each other but at different skill levels (junior and senior).

        Two other positions should be similar to each other, including in terms of skill level. 

        The remaining position should be different to the other four.

        Provide only the jsonl file contents, do not provide anything else or say anything else. 

        Everything after <jsonl> should be a jsonl_string ready to be processed by this Python code. 

        <code>
        list_of_dicts = [json.loads(line) for line in jsonl_string.strip().splitlines()]
        </code>

        ### OUTPUT ###
        <jsonl>
        """
    )

    jsonl_string_pds = (
        response_pd.content.replace("`", "").replace("json", "").replace("jsonl", "")
    )

    position_data = [json.loads(line) for line in jsonl_string_pds.strip().splitlines()]

    ut.write_to_csv(
        f"data/source/pds/{str(uuid.uuid4())}.csv",
        ["position_number", "job_title", "department", "level"],
        [
            [p["position_number"], p["job_title"], p["department"], p["level"]]
            for p in position_data
        ],
    )

    protected_characteristic_data = []
    suitability_data = []
    for p in position_data:
        create_pd_pdf(p)
        response_cvs = services.llm_hot.invoke(
            f"""
            ### BACKGROUND ###
            You are an expert career assistant. Your primary function is to analyze a job position description
            and generate data which could be found in the cvs of people applying for the position.

            From the position description provided below, identify and extract the key details.

            Structure your output as a single JSONL object.

            ### Instructions ###

            Carefully read the position description provided. 

            Identify the core duties, mandatory qualifications, and key skills.

            Compile a list of important keywords and phrases that an Applicant Tracking System (ATS) would
            likely scan for.

            Present all extracted information in the specified JSON format. Ensure that the
            key_responsibilities are phrased as action-oriented statements suitable for a CV.

            Position Description to Analyze:

            {p}

            Required Output Format:

            Provide your response as a single jsonl file only. Do not include any explanations or text outside
            of the jsonl content.

            This is an example of a row in the jsonl file.
            <output example>
            {{"suitability":"Y","personal_info":{{"full_name":"Jane Doe","job_title":"Senior Software Engineer","email":"jane.doe@email.com","phone":"+61 400 123 456","linkedin":"linkedin.com/in/janedoe","github":"github.com/janedoe","website":"janedoe.com","address":"123 Python Street, Melbourne, VIC 3000","summary":"A highly skilled and motivated senior software engineer with over 10 years of experience in developing scalable web applications. Proficient in Python, Django, and cloud technologies. Seeking to leverage my expertise to contribute to a dynamic engineering team."}},"work_experience":[{{"job_title":"Senior Software Engineer","company":"Tech Solutions Inc.","location":"Melbourne, VIC","start_date":"January 2020","end_date":"Present","responsibilities":["Led a team of 5 engineers in developing and maintaining a high-traffic e-commerce platform.","Architected and implemented a microservices-based backend using Python, Django, and Docker.","Improved application performance by 30% through database optimization and caching strategies.","Mentored junior developers and conducted code reviews to ensure code quality and best practices."]}},{{"job_title":"Software Engineer","company":"Digital Innovations Co.","location":"Sydney, NSW","start_date":"June 2015","end_date":"December 2019","responsibilities":["Developed RESTful APIs for a client-facing financial services application.","Collaborated with product managers and designers to translate requirements into technical solutions.","Wrote and maintained comprehensive unit and integration tests, achieving 95% code coverage."]}}],"education":[{{"degree":"Master of Information Technology","institution":"University of Melbourne","location":"Melbourne, VIC","graduation_year":"2015","details":"Specialization in Software Engineering. Thesis on machine learning applications."}},{{"degree":"Bachelor of Computer Science","institution":"Monash University","location":"Melbourne, VIC","graduation_year":"2013","details":"Graduated with High Distinction."}}],"skills":{{"programming_languages":["Python","JavaScript","SQL","Go"],"frameworks_libraries":["Django","Flask","React","Node.js","Pandas"],"databases":["PostgreSQL","MySQL","MongoDB","Redis"],"cloud_devops":["AWS","Docker","Kubernetes","CI/CD","Terraform"],"languages":[{{"language":"English","proficiency":"Native"}},{{"language":"Spanish","proficiency":"Intermediate"}}]}},"projects":[{{"name":"Personal Finance Tracker","description":"A web application built with Flask and PostgreSQL to track personal expenses and generate financial reports.","technologies":["Python","Flask","PostgreSQL","Chart.js"],"link":"github.com/janedoe/finance-tracker"}},{{"name":"Open Source Contribution: PyData","description":"Contributed to the open-source data analysis library by fixing bugs and adding new features for data visualization.","technologies":["Python","Pandas","Matplotlib"],"link":"github.com/pydata/pandas/pull/12345"}}],"certifications":[{{"name":"AWS Certified Solutions Architect - Associate","issuing_organization":"Amazon Web Services","date_issued":"August 2021"}}],"references":"Available upon request"}}. # noqa: E501
            </output example>

            ### Instructions ###
            You must now create a jsonl file with ten lines of data following the format of the above output
            example. Two of the lines should correspond to cvs which obviously meet the position criteria.
            Three of the lines should correspond to cvs that do not meet the position criteria. And five
            lines should correspond to cvs which most of the position description criteria but not all of it. 

            Use a diverse range of names and personalities for the people the cvs correspond to.

            You must include the suitability key. It takes on values Y or N respectively indicating whether
            the cv belongs to a candidate that is suitable for the position (Y) or not (N).

            YOU MUST MAKE SURE THAT WHAT YOU PROVIDE IN YOUR RESPONSE IS VALID CONTENT IN A JSONL FILE.

            <jsonl>
            """
        )

        jsonl_string_cvs = (
            response_cvs.content.replace("`", "")
            .replace("jsonl", "")
            .replace("json", "")
        )
        try:
            cv_data = [
                json.loads(line) for line in jsonl_string_cvs.strip().splitlines()
            ]

            for c in cv_data:
                id = str(uuid.uuid4())
                create_cv_pdf(c, p["position_number"], id)
                protected_characteristic_data.append(
                    generate_demographic_profiles(p["position_number"], id).values()
                )
                suitability_data.append([p["position_number"], id, c["suitability"]])
        except Exception as e:
            log.error(e)

    ut.write_to_csv(
        f"data/source/cvs/{str(uuid.uuid4())}.csv",
        generate_demographic_profiles("", "").keys(),
        protected_characteristic_data,
    )

    if SIMULATE_MANUAL_REVIEW:
        ut.write_to_csv(
            f"data/source/cvs/{str(uuid.uuid4())}-manual-review.csv",
            ["position_number", "application_id", "suitability_manual"],
            suitability_data,
        )
